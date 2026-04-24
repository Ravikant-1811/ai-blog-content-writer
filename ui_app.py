#!/usr/bin/env python3
"""
Basic web UI for AI Content Writer.

Run:
  export ANTHROPIC_API_KEY="your_key"
  python3 ui_app.py
"""

from __future__ import annotations

import os
import uuid
from datetime import timedelta
from pathlib import Path
from typing import Optional

from functools import wraps

from flask import (
    Flask,
    abort,
    flash,
    g,
    redirect,
    render_template,
    request,
    session,
    send_from_directory,
    url_for,
)
from werkzeug.security import check_password_hash, generate_password_hash

from claude_blog_writer import (
    DEFAULT_MODEL,
    build_http_session,
    call_claude,
    fetch_url_html,
    get_verify_bundle,
    html_to_text,
    save_pdf_output,
    validate_source_url,
)
from auth_store import create_user, fetch_user_by_email, fetch_user_by_id, init_db


app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", os.getenv("SECRET_KEY", "dev-change-me"))
app.permanent_session_lifetime = timedelta(days=7)
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
)
BASE_DIR = Path(__file__).resolve().parent
PUBLIC_ENDPOINTS = {"login", "register", "logout", "static"}

# Vercel serverless filesystem is read-only except /tmp.
if os.getenv("VERCEL"):
    OUTPUT_DIR = Path("/tmp/generated_outputs")
else:
    OUTPUT_DIR = BASE_DIR / "generated_outputs"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
init_db()


def login_required(view):
    @wraps(view)
    def wrapped_view(*args, **kwargs):
        if g.user is None:
            flash("Please sign in to access the writer.", "notice")
            return redirect(url_for("login"))
        return view(*args, **kwargs)

    return wrapped_view


@app.before_request
def load_current_user():
    user_id = session.get("user_id")
    user = fetch_user_by_id(user_id) if user_id else None
    if user_id and user is None:
        session.clear()
    g.user = user

    if request.endpoint and request.endpoint not in PUBLIC_ENDPOINTS and g.user is None:
        flash("Please sign in to access the writer.", "notice")
        return redirect(url_for("login", next=request.path))


@app.context_processor
def inject_current_user():
    return {"current_user": getattr(g, "user", None)}


def _safe_next_target(raw_target: str | None) -> str | None:
    if not raw_target:
        return None
    if raw_target.startswith("/") and not raw_target.startswith("//"):
        return raw_target
    return None


def save_docx_output(out_base: Path, content: str) -> Optional[Path]:
    try:
        from docx import Document
    except Exception:
        return None

    out_file = out_base.with_suffix(".docx")
    document = Document()
    for line in content.splitlines():
        if line.strip():
            document.add_paragraph(line)
        else:
            document.add_paragraph("")
    document.save(out_file)
    return out_file


@app.route("/login", methods=["GET", "POST"])
def login():
    if g.user is not None:
        return redirect(url_for("index"))

    error_msg = ""
    form = {"email": ""}
    next_target = _safe_next_target(request.args.get("next"))

    if request.method == "POST":
        form["email"] = (request.form.get("email") or "").strip()
        password = request.form.get("password") or ""
        next_target = _safe_next_target(request.form.get("next")) or next_target

        user = fetch_user_by_email(form["email"])
        if user is None or not check_password_hash(user["password_hash"], password):
            error_msg = "Invalid email or password."
        else:
            session.clear()
            session.permanent = True
            session["user_id"] = user["id"]
            flash(f"Welcome back, {user['name']}!", "info")
            return redirect(next_target or url_for("index"))

    return render_template(
        "auth.html",
        page_title="Sign in",
        heading="Sign in",
        subheading="Use your account to access the content writer.",
        mode="login",
        form=form,
        error_msg=error_msg,
        next_target=next_target,
    )


@app.route("/register", methods=["GET", "POST"])
def register():
    if g.user is not None:
        return redirect(url_for("index"))

    error_msg = ""
    form = {"name": "", "email": ""}
    next_target = _safe_next_target(request.args.get("next"))

    if request.method == "POST":
        form["name"] = (request.form.get("name") or "").strip()
        form["email"] = (request.form.get("email") or "").strip()
        password = request.form.get("password") or ""
        confirm_password = request.form.get("confirm_password") or ""
        next_target = _safe_next_target(request.form.get("next")) or next_target

        if not form["name"] or not form["email"] or not password:
            error_msg = "Name, email, and password are required."
        elif len(password) < 8:
            error_msg = "Password must be at least 8 characters long."
        elif password != confirm_password:
            error_msg = "Passwords do not match."
        elif fetch_user_by_email(form["email"]) is not None:
            error_msg = "An account with that email already exists."
        else:
            password_hash = generate_password_hash(password)
            user = create_user(form["name"], form["email"], password_hash)
            session.clear()
            session.permanent = True
            session["user_id"] = user["id"]
            flash("Account created successfully.", "info")
            return redirect(next_target or url_for("index"))

    return render_template(
        "auth.html",
        page_title="Create account",
        heading="Create account",
        subheading="Set up the workspace account for content generation.",
        mode="register",
        form=form,
        error_msg=error_msg,
        next_target=next_target,
    )


@app.get("/logout")
def logout():
    session.clear()
    flash("You have been signed out.", "info")
    return redirect(url_for("login"))


@app.get("/download/<path:filename>")
@login_required
def download_file(filename: str):
    safe_name = Path(filename).name
    if safe_name != filename:
        abort(404)
    target = OUTPUT_DIR / safe_name
    if not target.exists() or not target.is_file():
        abort(404)
    return send_from_directory(OUTPUT_DIR, safe_name, as_attachment=True)


@app.route("/", methods=["GET", "POST"])
@login_required
def index():
    data = {
        "url": "",
        "prompt": "",
        "model": DEFAULT_MODEL,
        "max_tokens": "2500",
        "temperature": "0.4",
    }
    result_text = ""
    error_msg = ""
    info_msg = ""
    download_files: dict[str, str] = {}
    backend_api_key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    show_api_key_input = not bool(backend_api_key)

    if request.method == "POST":
        data["url"] = (request.form.get("url") or "").strip()
        data["prompt"] = (request.form.get("prompt") or "").strip()
        data["model"] = (request.form.get("model") or DEFAULT_MODEL).strip()
        data["max_tokens"] = (request.form.get("max_tokens") or "2500").strip()
        data["temperature"] = (request.form.get("temperature") or "0.4").strip()

        api_key = backend_api_key or (request.form.get("api_key") or "").strip()
        if not api_key:
            error_msg = (
                "Please enter your Anthropic API key, or configure "
                "ANTHROPIC_API_KEY on the backend."
            )
            return render_template(
                "index.html",
                data=data,
                result_text=result_text,
                error_msg=error_msg,
                info_msg=info_msg,
                download_files=download_files,
                show_api_key_input=show_api_key_input,
            )

        if not data["url"] or not data["prompt"]:
            error_msg = "URL and Prompt are required."
            return render_template(
                "index.html",
                data=data,
                result_text=result_text,
                error_msg=error_msg,
                info_msg=info_msg,
                download_files=download_files,
                show_api_key_input=show_api_key_input,
            )

        try:
            max_tokens = int(data["max_tokens"])
            temperature = float(data["temperature"])
        except ValueError:
            error_msg = "Invalid max_tokens or temperature value."
            return render_template(
                "index.html",
                data=data,
                result_text=result_text,
                error_msg=error_msg,
                info_msg=info_msg,
                download_files=download_files,
                show_api_key_input=show_api_key_input,
            )

        try:
            validate_source_url(data["url"])
            session = build_http_session()
            verify_bundle = get_verify_bundle()

            html = fetch_url_html(session, data["url"], verify=verify_bundle)
            page_title, page_text = html_to_text(html)

            result_text, was_truncated = call_claude(
                session=session,
                api_key=api_key,
                model=data["model"],
                prompt=data["prompt"],
                source_url=data["url"],
                page_title=page_title,
                page_text=page_text,
                verify=verify_bundle,
                max_tokens=max_tokens,
                temperature=temperature,
            )

            file_id = uuid.uuid4().hex
            out_base = OUTPUT_DIR / f"ai_content_{file_id}"

            md_file = out_base.with_suffix(".md")
            md_file.write_text(result_text, encoding="utf-8")
            download_files["Markdown"] = md_file.name

            pdf_file = save_pdf_output(out_base, result_text)
            if pdf_file:
                download_files["PDF"] = pdf_file.name
            else:
                info_msg = "PDF export unavailable (install reportlab)."

            docx_file = save_docx_output(out_base, result_text)
            if docx_file:
                download_files["Word"] = docx_file.name
            else:
                if info_msg:
                    info_msg += " "
                info_msg += "Word export unavailable (install python-docx)."

            if was_truncated:
                if info_msg:
                    info_msg += " "
                info_msg += "Successfully done."
            elif not info_msg:
                info_msg = "Successfully done."

        except Exception as exc:
            error_msg = str(exc)

    return render_template(
        "index.html",
        data=data,
        result_text=result_text,
        error_msg=error_msg,
        info_msg=info_msg,
        download_files=download_files,
        show_api_key_input=show_api_key_input,
    )


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5050, debug=True)
